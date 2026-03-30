"""
Utility module for exporting group data to DOCX format
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def get_export_translations(language='en'):
    """
    Get translation strings for export based on language
    
    Args:
        language: 'en' for English or 'bg' for Bulgarian
    
    Returns:
        Dictionary with all export-related translation strings
    """
    translations = {
        'en': {
            'group_report': 'Group Report',
            'generated_on': 'Generated on',
            'group_overview': 'Group Overview',
            'group_name': 'Group Name',
            'description': 'Description',
            'no_description': 'No description',
            'created_by': 'Created By',
            'created_date': 'Created Date',
            'total_members': 'Total Members',
            'total_messages': 'Total Messages',
            'members': 'Members',
            'total_members_label': 'Total members',
            'number': '#',
            'name': 'Name',
            'username': 'Username',
            'role': 'Role',
            'joined': 'Joined',
            'messages': 'Messages',
            'total_messages_label': 'Total messages',
            'no_messages': 'No messages yet.',
            'date_format': '%B %d, %Y',
            'date_time_format': '%B %d, %Y at %I:%M %p',
            'message_date_format': '%b %d, %Y %I:%M %p',
        },
        'bg': {
            'group_report': 'Доклад за група',
            'generated_on': 'Генериран на',
            'group_overview': 'Преглед на група',
            'group_name': 'Име на група',
            'description': 'Описание',
            'no_description': 'Нямаме описание',
            'created_by': 'Създадено от',
            'created_date': 'Дата на създаване',
            'total_members': 'Общо членове',
            'total_messages': 'Общо съобщения',
            'members': 'Членове',
            'total_members_label': 'Общо членове',
            'number': '#',
            'name': 'Име',
            'username': 'Потребителско име',
            'role': 'Роля',
            'joined': 'Присъединен/а',
            'messages': 'Съобщения',
            'total_messages_label': 'Общо съобщения',
            'no_messages': 'Все още няма съобщения.',
            'date_format': '%d.%m.%Y',
            'date_time_format': '%d.%m.%Y в %H:%M',
            'message_date_format': '%d.%m.%Y %H:%M',
        }
    }
    return translations.get(language, translations['en'])


def generate_group_export(group, members, messages, language='en'):
    """
    Generate a DOCX document with group information
    
    Args:
        group: Group object
        members: List of GroupMember objects with related user data
        messages: List of Message objects for the group
        language: 'en' for English or 'bg' for Bulgarian
    
    Returns:
        Path to the generated DOCX file
    """
    t = get_export_translations(language)
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'{t["group_report"]}: {group.name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add generation timestamp
    timestamp = doc.add_paragraph()
    timestamp.add_run(f'{t["generated_on"]} {datetime.now().strftime(t["date_time_format"])}').italic = True
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # ===== GROUP OVERVIEW SECTION =====
    doc.add_heading(t['group_overview'], 1)
    
    overview_table = doc.add_table(rows=6, cols=2)
    overview_table.style = 'Light Grid Accent 1'
    
    # Set column widths
    overview_table.autofit = False
    for row in overview_table.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(3.5)
    
    # Fill overview data
    overview_data = [
        (t['group_name'], group.name),
        (t['description'], group.description or t['no_description']),
        (t['created_by'], group.created_by.get_full_name() or group.created_by.username),
        (t['created_date'], group.created_at.strftime(t['date_format']) if group.created_at else 'N/A'),
        (t['total_members'], str(len(members))),
        (t['total_messages'], str(len(messages))),
    ]
    
    for idx, (label, value) in enumerate(overview_data):
        row = overview_table.rows[idx]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        # Bold the labels
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # ===== MEMBERS SECTION =====
    doc.add_heading(t['members'], 1)
    doc.add_paragraph(f'{t["total_members_label"]}: {len(members)}')
    
    # Create members table
    members_table = doc.add_table(rows=len(members) + 1, cols=5)
    members_table.style = 'Light Grid Accent 1'
    
    # Set column widths
    members_table.autofit = False
    for row in members_table.rows:
        row.cells[0].width = Inches(1.2)
        row.cells[1].width = Inches(1.5)
        row.cells[2].width = Inches(1.5)
        row.cells[3].width = Inches(1.2)
        row.cells[4].width = Inches(1.1)
    
    # Header row
    header_cells = members_table.rows[0].cells
    headers = [t['number'], t['name'], t['username'], t['role'], t['joined']]
    for idx, header_text in enumerate(headers):
        header_cells[idx].text = header_text
        # Bold header
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        # Header background color (dark)
        shading_elm = header_cells[idx]._element
        shading_elm.get_or_add_tcPr().append(
            doc._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
        )
    
    # Member rows
    for idx, member in enumerate(members, 1):
        row = members_table.rows[idx]
        row.cells[0].text = str(idx)
        row.cells[1].text = member.user.get_full_name() or member.user.username
        row.cells[2].text = member.user.username
        row.cells[3].text = member.role.replace('_', ' ').title()
        row.cells[4].text = member.joined_at.strftime(t['date_format']) if member.joined_at else 'N/A'
    
    doc.add_paragraph()  # Spacing
    
    # ===== MESSAGES SECTION =====
    doc.add_heading(t['messages'], 1)
    doc.add_paragraph(f'{t["total_messages_label"]}: {len(messages)}')
    
    if messages:
        for msg in messages:
            # Message sender and timestamp
            msg_header = doc.add_paragraph()
            msg_sender_run = msg_header.add_run(msg.sender.get_full_name() or msg.sender.username)
            msg_sender_run.bold = True
            msg_header.add_run(f' @{msg.sender.username} • {msg.created_at.strftime(t["message_date_format"])}').italic = True
            
            # Message content
            msg_content = doc.add_paragraph(msg.content, style='List Bullet')
            msg_content.paragraph_format.left_indent = Inches(0.5)
    else:
        doc.add_paragraph(t['no_messages'])
    
    # Create export filename
    timestamp_str = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f'Group_{group.id}_{group.name.replace(" ", "_")}_{timestamp_str}.docx'
    filepath = os.path.join(
        os.path.dirname(os.path.dirname(__file__)),
        'exports',
        filename
    )
    
    # Ensure exports directory exists
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    
    # Save document
    doc.save(filepath)
    
    return filepath, filename
